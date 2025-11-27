"""Ingestion Tools - Optimized document processing and database tracking for DeepAgents RAG."""
import json
import sqlite3
import datetime
from pathlib import Path
from typing import Dict, Any, List
from langchain_core.tools import tool
from langchain_text_splitters import RecursiveCharacterTextSplitter
# Assuming relative imports for configuration and database models
from ..config.env_config import EnvConfig 

# --- Optional Document Extraction Libraries ---
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from PyPDF2 import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ============================================================================
# DOCUMENT EXTRACTION HELPERS (Internal functions, no @tool decorator)
# ============================================================================

def extract_pdf_text_pdfplumber(pdf_path: str, include_tables: bool = True) -> Dict[str, Any]:
    """
    Internal helper: Extracts text from PDF using pdfplumber, prioritizing structured data like tables.
    """
    if not HAS_PDFPLUMBER:
        return {"success": False, "error": "pdfplumber not installed."}
    
    try:
        text_content = ""
        metadata = {"pages": 0, "tables_found": 0}
        with pdfplumber.open(pdf_path) as pdf:
            metadata["pages"] = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages):
                text_content += f"\n--- Page {page_num + 1} ---\n"
                text_content += page.extract_text() or ""
                
                if include_tables:
                    tables = page.extract_tables()
                    if tables:
                        metadata["tables_found"] += len(tables)
                        text_content += "\n**Tables on this page:**\n"
                        for table_idx, table in enumerate(tables):
                            text_content += f"\n**Table {table_idx + 1}:**\n"
                            for row in table:
                                text_content += " | ".join(str(cell or "") for cell in row) + "\n"
        
        return {"success": True, "text": text_content, "metadata": metadata, "method": "pdfplumber"}
    except Exception as e:
        return {"success": False, "error": f"pdfplumber extraction failed: {str(e)}"}


def extract_pdf_text_pypdf(pdf_path: str) -> Dict[str, Any]:
    """
    Internal helper: Extracts text from PDF using PyPDF2 as a reliable fallback.
    """
    if not HAS_PYPDF:
        return {"success": False, "error": "PyPDF2 not installed."}
    
    try:
        text_content = ""
        reader = PdfReader(pdf_path)
        pages_count = len(reader.pages)
        
        for page_num, page in enumerate(reader.pages):
            text_content += f"\n--- Page {page_num + 1} ---\n"
            text_content += page.extract_text() or ""
        
        return {"success": True, "text": text_content, "metadata": {"pages": pages_count}, "method": "PyPDF2"}
    except Exception as e:
        return {"success": False, "error": f"PyPDF2 extraction failed: {str(e)}"}


def extract_text_file(file_path: str) -> Dict[str, Any]:
    """Internal helper: Extracts raw text from standard text files (.txt, .md)."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return {
            "success": True,
            "text": text,
            "metadata": {"lines": len(text.split('\n'))},
            "method": "text_file"
        }
    except Exception as e:
        return {"success": False, "error": f"Text extraction failed: {str(e)}"}


def extract_word_file(file_path: str) -> Dict[str, Any]:
    """Internal helper: Extracts text from Microsoft Word documents (.docx)."""
    if not HAS_DOCX:
        return {"success": False, "error": "python-docx not installed."}
    
    try:
        doc = DocxDocument(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        table_count = len(doc.tables)
        for table in doc.tables:
            text_parts.append("\n**Table:**\n")
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                text_parts.append(" | ".join(cells))
        
        text = "\n".join(text_parts)
        return {
            "success": True,
            "text": text,
            "metadata": {"paragraphs": len(doc.paragraphs), "tables": table_count},
            "method": "word_file"
        }
    except Exception as e:
        return {"success": False, "error": f"Word extraction failed: {str(e)}"}


# ============================================================================
# AGENT TOOL FUNCTIONS (@tool decorators applied)
# ============================================================================

@tool
def ingest_documents_from_path_tool(path: str, doc_id_prefix: str = "doc", 
                                     file_type: str = "auto", recursive: bool = True) -> str:
    """
    **BATCH DISCOVERY TOOL.** Scans a local file path or folder (recursively) for 
    supported documents (PDF, TXT, DOCX). It performs path validation and file discovery, 
    but **does NOT perform extraction**. It returns a list of files to be processed by the agent.
    
    Args:
        path (str): File path or folder path to scan.
        doc_id_prefix (str): Prefix for auto-generated doc_ids.
        file_type (str): File type filter - "auto", "pdf", "text", "word", or "all".
        recursive (bool): Search subdirectories recursively.
    
    Returns:
        str: JSON with 'success' status, total count, scan path, and a list of 
             `discovered_documents` (each containing `file_path`, `file_type`, and `doc_id`).
             Example Success: {"success": true, "total_discovered": 5, "discovered_documents": [...]}
             Example Failure (Graceful): {"success": false, "error": "Path not found: ...", "discovered_documents": []}
    """
    try:
        path_obj = Path(path)
        
        if not path_obj.exists():
            return json.dumps({
                "success": False,
                "error": f"Path not found: {path}",
                "total_discovered": 0,
                "discovered_documents": []
            })
        
        supported_exts = {
            '.pdf': 'pdf', '.txt': 'text', '.md': 'text', '.markdown': 'text',
            '.docx': 'word', '.doc': 'word'
        }
        
        files_to_process = []
        
        # Find all files
        if path_obj.is_file():
            files_to_process = [path_obj]
        elif path_obj.is_dir():
            if recursive:
                files_to_process = [f for f in path_obj.rglob('*') if f.is_file()]
            else:
                files_to_process = [f for f in path_obj.glob('*') if f.is_file()]
        
        # Filter by extension
        filtered_files = []
        for f in files_to_process:
            ext = f.suffix.lower()
            detected_type = supported_exts.get(ext)
            
            if file_type == "all" or detected_type == file_type or (file_type == "auto" and detected_type):
                filtered_files.append(f)
        
        if not filtered_files:
            return json.dumps({
                "success": False,
                "error": f"No documents matching type '{file_type}' found in {path}",
                "discovered_documents": []
            })
        
        # Prepare discovery results
        discovered = []
        for file_path in filtered_files:
            file_ext = file_path.suffix.lower()
            file_type_detected = supported_exts.get(file_ext, 'unknown')
            
            from datetime import datetime as dt
            doc_id = f"{doc_id_prefix}_{file_path.stem}_{dt.now().strftime('%Y%m%d_%H%M%S')}"
            
            discovered.append({
                "file_path": str(file_path),
                "file_name": file_path.name,
                "file_type": file_type_detected,
                "doc_id": doc_id,
                "size_kb": file_path.stat().st_size / 1024
            })
        
        return json.dumps({
            "success": True,
            "total_discovered": len(discovered),
            "path_scanned": str(path),
            "recursive": recursive,
            "discovered_documents": discovered
        })
        
    except Exception as e:
        import traceback
        return json.dumps({
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc(),
            "discovered_documents": []
        })

@tool
def extract_document_text_tool(file_path: str, file_type: str = "auto") -> str:
    """
    **STAGE 1: DOCUMENT PROCESSING.** Extracts raw, structured text content 
    from various document types (PDF, TXT, DOCX) using internal specialized parsers. 
    This tool initiates the transformation of a file into text.

    Args:
        file_path (str): The local file path to the document.
        file_type (str): Explicitly set the file type ("pdf", "text", "word") 
                         or use "auto" to detect from the file extension.

    Returns:
        str: JSON string containing the 'success' status, the raw extracted 'text', 
             extraction 'metadata', and the 'method' used. 
             Example Success: {"success": true, "text": "...", "metadata": {...}, "method": "..."}
    """
    file_path = str(file_path)
    
    if not Path(file_path).exists():
        return json.dumps({"success": False, "error": f"File not found: {file_path}"})
    
    # Detect file type if auto
    if file_type == "auto":
        suffix = Path(file_path).suffix.lower()
        if suffix == '.pdf': file_type = 'pdf'
        elif suffix in ['.txt', '.md', '.markdown']: file_type = 'text'
        elif suffix in ['.docx', '.doc']: file_type = 'word'
        else: return json.dumps({"success": False, "error": f"Unsupported file type: {suffix}"})
    
    if file_type == 'pdf':
        result = extract_pdf_text_pdfplumber(file_path)
        if not result["success"]:
            # Fallback to PyPDF2 if pdfplumber fails
            result = extract_pdf_text_pypdf(file_path)
        return json.dumps(result)
    
    elif file_type == 'text':
        return json.dumps(extract_text_file(file_path))
    
    elif file_type == 'word':
        return json.dumps(extract_word_file(file_path))
    
    else:
        return json.dumps({"success": False, "error": f"Unknown file type: {file_type}"})

@tool
def extract_metadata_tool(text: str, llm_service) -> str:
    """
    **STAGE 2: METADATA ENRICHMENT.** Extracts high-level, semantic metadata 
    (title, summary, type) from the raw document text using the **LLM's structured output capability**. 
    This metadata is critical for RAG filtering and search relevance.
    
    Args:
        text (str): The full or truncated document text for analysis.
        llm_service: Service object providing a .generate_json(prompt) method.

    Returns:
        str: JSON string with 'success' and the extracted 'metadata' dictionary.
             Example Success: {"success": true, "metadata": {"title": "...", "summary": "...", "doc_type": "..."}}
    """
    try:
        prompt = f"""Extract metadata from this document:
        
Document (first 2000 chars):
{text[:2000]}

Return JSON with: title, summary (2-3 sentences), keywords (5-10 list), topics (list), doc_type (manual|policy|technical_doc|report|incident|table_ingest).

Example: {{"title": "...", "summary": "...", "keywords": [...], "topics": [...], "doc_type": "..."}}"""
        
        result = llm_service.generate_json(prompt)
        parsed = result if isinstance(result, dict) else json.loads(result)
        
        return json.dumps({"success": True, "metadata": parsed})
        
    except Exception as e:
        # Robust fallback on any failure
        fallback_metadata = {
            "title": "Document",
            "summary": "Unable to extract detailed metadata.",
            "keywords": ["document", "metadata_failure"],
            "topics": ["unknown"],
            "doc_type": "report"
        }
        return json.dumps({"success": True, "metadata": fallback_metadata, "error_detail": str(e)})


@tool
def chunk_document_tool(text: str, doc_id: str, strategy: str = "recursive", 
                        chunk_size: int = 500, overlap: int = 50) -> str:
    """
    **STAGE 3: CHUNKING.** Splits the long document text into smaller, 
    semantically coherent chunks using a `RecursiveCharacterTextSplitter`.
    
    Args:
        text (str): The document content to be chunked.
        doc_id (str): The unique ID of the source document/file.
        strategy (str): The splitting strategy (default 'recursive').
        chunk_size (int): Max number of characters per chunk.
        overlap (int): Overlap in characters between adjacent chunks for continuity.

    Returns:
        str: JSON string with 'success', 'doc_id', 'num_chunks', and a list of 'chunks'.
             Example Success: {"success": true, "doc_id": "...", "num_chunks": 10, "chunks": [...]}
    """
    try:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            # Optimized separators for Markdown/general text structure
            separators=["\n\n##", "\n\n", "\n", ". ", " ", ""]
        )
        
        chunks = splitter.split_text(text)
        result = [
            {
                "chunk_id": f"{doc_id}_chunk_{i}",
                "text": chunk,
                "strategy": strategy,
                "size": len(chunk),
                "index": i
            }
            for i, chunk in enumerate(chunks)
        ]
        
        return json.dumps({
            "success": True,
            "doc_id": doc_id,
            "num_chunks": len(result),
            "chunks": result
        })
        
    except Exception as e:
        return json.dumps({"success": False, "error": str(e)})


@tool
def save_to_vectordb_tool(chunks: str, doc_id: str, llm_service, vectordb_service, 
                          metadata: str = None, rbac_namespace: str = "general") -> str:
    """
    **STAGE 4: EMBEDDING & PERSISTENCE.** Generates embeddings for chunks, 
    stores them in the **Vector Database (VDB)** for similarity search, 
    and persists metadata into the local **SQLite tracking DB** (Document/Chunk Metadata Models).

    Args:
        chunks (str): JSON string returned by `chunk_document_tool`.
        doc_id (str): Unique ID for the source document/file.
        llm_service: Service object for .generate_embedding(text).
        vectordb_service: Service object for VDB .collection.add(...).
        metadata (str): JSON string of document-level metadata (from `extract_metadata_tool`).
        rbac_namespace (str): Namespace/Collection name used for Access Control filtering.

    Returns:
        str: JSON with 'success', 'doc_id', 'chunks_saved', and VDB details.
             Example Success: {"success": true, "doc_id": "...", "chunks_saved": 10, "rbac_namespace": "..."}
    """
    try:
        chunks_data = json.loads(chunks) if isinstance(chunks, str) else chunks
        chunk_list = chunks_data.get('chunks', [])
        
        if not chunks_data.get('success', False) or not chunk_list:
             return json.dumps({"success": False, "error": "Invalid or empty chunks list provided."})
        
        # 1. Parse Metadata
        doc_metadata = {}
        if metadata:
            meta_data = json.loads(metadata) if isinstance(metadata, str) else {}
            doc_metadata = meta_data.get('metadata', {})

        # 2. Prepare Data Structures
        chunk_ids, embeddings, texts, metadatas = [], [], [], []
        
        cleaned_doc_metadata = {
            "doc_id": doc_id,
            "rbac_namespace": rbac_namespace,
            "ingestion_date": datetime.datetime.now().isoformat(),
            **{k: (json.dumps(v) if isinstance(v, (list, dict)) else str(v)) 
               for k, v in doc_metadata.items()}
        }
        
        # 3. Process Chunks (Generate Embeddings & Append)
        for chunk in chunk_list:
            chunk_text = chunk.get('text', '').strip()
            if not chunk_text: continue
            
            chunk_id = f"{doc_id}_chunk_{chunk.get('index', 0)}"
            embedding = llm_service.generate_embedding(chunk_text)
            
            chunk_ids.append(chunk_id)
            embeddings.append(embedding)
            texts.append(chunk_text)
            
            metadatas.append({
                "chunk_index": chunk.get('index', 0),
                **cleaned_doc_metadata
            })
        
        if not chunk_ids:
            return json.dumps({"success": False, "error": "No valid chunk text found after processing."})
        
        # 4. Add to Vector DB (ChromaDB)
        vectordb_service.collection.add(
            ids=chunk_ids,
            documents=texts,
            embeddings=embeddings,
            metadatas=metadatas
        )
        
        # 5. Save to SQLite optimized schema (for audit and healing)
        try:
            # Assumed imports for models (must exist in the relative path)
            from ...database.models.document_metadata_model import DocumentMetadataModel
            from ...database.models.chunk_embedding_data_model import ChunkEmbeddingDataModel
            
            # Models now handle their own connection via get_connection()
            doc_model = DocumentMetadataModel()
            doc_model.create(
                doc_id=doc_id,
                title=doc_metadata.get('title', f'Document {doc_id}'),
                author=doc_metadata.get('author', 'Unknown'),
                source=doc_metadata.get('source', 'ingestion_tool'),
                summary=doc_metadata.get('summary', ''),
                rbac_namespace=rbac_namespace,
                chunk_strategy="recursive_splitter",
                chunk_size_char=500,
                overlap_char=50,
                metadata_json=json.dumps(doc_metadata)
            )
            
            chunk_model = ChunkEmbeddingDataModel()
            for i, chunk_id in enumerate(chunk_ids):
                chunk_model.create(
                    chunk_id=chunk_id,
                    doc_id=doc_id,
                    embedding_model=llm_service.provider if hasattr(llm_service, 'provider') else 'default_embedder',
                    embedding_version="1.0",
                    quality_score=0.8, 
                    reindex_count=0,
                    healing_suggestions=json.dumps({})
                )
            
            doc_model.close()
            chunk_model.close()
            
        except Exception as e:
            # Log SQLite failure but continue if VDB write succeeded
            print(f"[ERROR] SQLite metadata save failed: {e}")
        
        return json.dumps({
            "success": True,
            "doc_id": doc_id,
            "chunks_saved": len(chunk_ids),
            "rbac_namespace": rbac_namespace,
        })
        
    except Exception as e:
        return json.dumps({"success": False, "error": str(e)})


@tool
def update_metadata_tracking_tool(doc_id: str, source_path: str, rbac_namespace: str, 
                                 metadata: str, chunks_saved: int, is_table: bool = False) -> str:
    """
    **STAGE 5: FINAL AUDIT LOG.** Updates the central high-level **DocumentTrackingModel**
    in SQLite to finalize the ingestion record, marking the document as completed 
    with a traceable audit trail.
    
    Args:
        doc_id (str): Unique identifier for the document or table.
        source_path (str): Original file path or DB/table reference.
        rbac_namespace (str): The domain/namespace used for VDB.
        metadata (str): JSON string of document-level metadata (from LLM extraction).
        chunks_saved (int): Number of embeddings successfully stored in VDB.
        is_table (bool): True if the source was a database table.

    Returns:
        str: JSON with 'success' status. Example Success: {"success": true}
    """
    try:
        rag_db_path = EnvConfig.get_db_path()
        rag_conn = sqlite3.connect(rag_db_path)
        
        # Attempted import for model - gracefully handle if unavailable
        doc_metadata_dict = json.loads(metadata) if isinstance(metadata, str) else {}
        
        try:
            from ...database.models.document_tracking_model import DocumentTrackingModel 
            doc_model = DocumentTrackingModel(rag_conn)
            
            doc_model.insert({
                'document_id': doc_id,
                'source_path': source_path,
                'rbac_namespace': rbac_namespace,
                'doc_type': doc_metadata_dict.get('doc_type', 'unknown'),
                'chunks_saved': chunks_saved,
                'is_table': 1 if is_table else 0,
                'ingestion_date': datetime.datetime.now().isoformat(),
                'ingestion_status': 'COMPLETED',
                'metadata_tags': json.dumps(doc_metadata_dict)
            })
            
            rag_conn.commit()
            doc_model.close()
        except ImportError:
            # DocumentTrackingModel not available - log message but continue
            print(f"[WARNING] DocumentTrackingModel not available for doc_id={doc_id}")
        
        rag_conn.close()
        return json.dumps({"success": True})
        
    except Exception as e:
        return json.dumps({"success": False, "error": f"Metadata tracking failed: {str(e)}"})
    
    
@tool
def ingest_sqlite_table_tool(table_name: str, doc_id: str, rbac_namespace: str,
                              text_columns: list, metadata_columns: list = None,
                              db_path: str = None, llm_service=None, vectordb_service=None,
                              chunk_size: int = 512, chunk_overlap: int = 50,
                              where_clause: str = None) -> str:
    """
    **COMPLETE PIPELINE: TABLE INGESTION.** Ingests a single SQLite table into the RAG vector database. 
    It converts rows to structured text, chunks, embeds, and stores them in VDB.
    
    Args:
        table_name (str): Name of SQLite table to ingest.
        doc_id (str): Unique ingestion ID (used as the overall document ID).
        rbac_namespace (str): Domain/namespace for RBAC filtering.
        text_columns (list): List of column names to combine as searchable text.
        metadata_columns (list): List of column names to attach as metadata (optional).
        db_path (str): Path to SQLite database (uses config if None).
        llm_service: LLM service for embeddings.
        vectordb_service: Vector DB service.
        chunk_size (int): Max chunk size in characters/tokens.
        chunk_overlap (int): Overlap between chunks.
        where_clause (str): Optional SQL WHERE clause for filtering records.
    
    Returns:
        str: JSON with ingestion status and statistics.
             Example Success: {"success": true, "table_name": "...", "records_processed": 100, "total_chunks_saved": 150}
    """
    try:
        # 1. Load KB database path from data_sources.json configuration
        # The KB database (incident_iq.db) is read-only source for knowledge data
        # The RAG database (rag.db) is write-only for metadata and embeddings
        config_path = Path(__file__).parent.parent / "config" / "data_sources.json"
        try:
            with open(config_path, 'r') as f:
                config = json.load(f)
            kb_db_path = config.get("data_sources", {}).get("sqlite", {}).get("connection_string_env", "incident_iq.db")
        except:
            kb_db_path = "incident_iq.db"  # Fallback to default
        
        # Setup & Fetch Records from KNOWLEDGE BASE DATABASE (not RAG DB)
        conn = sqlite3.connect(kb_db_path)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        query = f"SELECT * FROM {table_name}"
        if where_clause:
            query += f" WHERE {where_clause}"
        cursor.execute(query)
        rows = cursor.fetchall()
        conn.close()
        
        if not rows:
            return json.dumps({"success": False, "error": f"No records found in table '{table_name}'"})
        
        # 2. Process Records (Convert to Structured Documents)
        documents_to_chunk = []
        for row_idx, row in enumerate(rows):
            row_dict = dict(row)
            
            # Create a structured Markdown-like text for high-quality embedding
            text_parts = [f"--- Table Record: {table_name} (Index: {row_idx}) ---"]
            for col in text_columns:
                value = row_dict.get(col)
                if value is not None:
                    text_parts.append(f"**{col.replace('_', ' ').title()}:** {value}")
            
            document_text = "\n".join(text_parts)
            
            metadata = {
                "source_table": table_name,
                "rbac_namespace": rbac_namespace,
                "doc_type": "table_record",
                "source_record_index": row_idx 
            }
            if metadata_columns:
                for col in metadata_columns:
                    if col in row_dict:
                        metadata[col] = row_dict[col]

            documents_to_chunk.append({"text": document_text, "metadata": metadata})

        # 3. Chunking, Embedding, and Storing
        splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        all_chunk_ids, all_texts, all_embeddings, all_metadatas = [], [], [], []
        
        for doc_idx, doc in enumerate(documents_to_chunk):
            chunks = splitter.split_text(doc["text"])
            
            for chunk_idx, chunk_text in enumerate(chunks):
                chunk_id = f"{doc_id}_{doc_idx}_{chunk_idx}" 
                
                embedding = llm_service.generate_embedding(chunk_text)
                
                chunk_metadata = {
                    "chunk_index": chunk_idx,
                    "doc_id": doc_id,
                    **doc["metadata"]
                }
                
                all_chunk_ids.append(chunk_id)
                all_texts.append(chunk_text)
                all_embeddings.append(embedding)
                all_metadatas.append(chunk_metadata)

        # 4. Save to Vector DB
        vectordb_service.collection.add(
            ids=all_chunk_ids,
            documents=all_texts,
            embeddings=all_embeddings,
            metadatas=all_metadatas
        )

        return json.dumps({
            "success": True,
            "doc_id": doc_id,
            "table_name": table_name,
            "records_processed": len(rows),
            "total_chunks_saved": len(all_chunk_ids),
            "rbac_namespace": rbac_namespace
        })
        
    except Exception as e:
        import traceback
        return json.dumps({
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        })
    
@tool
def record_agent_memory_tool(agent_name: str, memory_key: str, memory_value: str, 
                              memory_type: str = "context") -> str:
    """
    **AGENT UTILITY.** Stores agent memory, context, logs, or strategic decisions 
    into the `AgentMemoryModel` in SQLite for future self-reflection and debugging.
    
    Args:
        agent_name (str): The name of the agent logging the memory (e.g., 'ingestion-agent').
        memory_key (str): A unique key to identify the memory content (e.g., 'ingestion_status_report').
        memory_value (str): The content to be stored (often a JSON or summary string).
        memory_type (str): Type of memory ("context", "log", "decision", "performance").

    Returns:
        str: JSON with 'success' status. Example Success: {"success": true}
    """
    try:
        rag_db_path = EnvConfig.get_db_path()
        rag_conn = sqlite3.connect(rag_db_path)
        
        # Attempted import for model - gracefully handle if unavailable
        try:
            from ...database.models.agent_memory_model import AgentMemoryModel
            mem_model = AgentMemoryModel(rag_conn)
            
            mem_model.insert({
                'agent_id': agent_name,
                'memory_type': memory_type,
                'content': memory_value,
                'importance_score': 0.8,
                'created_at': datetime.datetime.now().isoformat()
            })
            
            rag_conn.commit()
            mem_model.close()
        except ImportError:
            # AgentMemoryModel not available - log message but continue
            print(f"[WARNING] AgentMemoryModel not available for agent={agent_name}")
        
        rag_conn.close()
        return json.dumps({"success": True})
        
    except Exception as e:
        return json.dumps({"success": False, "error": str(e)})
